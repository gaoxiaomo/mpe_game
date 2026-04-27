from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
TEMPLATE = ROOT / "midterm_template.docx"
OUTPUT_DIR = ROOT / "outputs" / "midterm_report_final"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

DOCX_OUT = OUTPUT_DIR / "midterm_report_draft.docx"
MD_OUT = OUTPUT_DIR / "midterm_report_draft.md"
QA_OUT = OUTPUT_DIR / "midterm_possible_questions.md"

TITLE = "基于强化学习的多智能体系统智能追逃博弈研究"
STUDENT_ID = "22371242"
SCHOOL = "计算机学院"
MAJOR = "计算机科学与技术"
STUDENT = "高悠然"
ADVISOR = "吴际"
DATE_TEXT = "2026年4月11日"


STD_RESULTS = [
    ["3v1", "52.05", "47.15", "275.8", "135.3", "1734.30", "1693.24"],
    ["3v3", "33.75", "33.75", "1361.3", "1361.3", "548.18", "548.18"],
    ["5v3", "36.80", "36.80", "185.0", "167.1", "553.03", "552.20"],
    ["6v3", "36.90", "37.75", "112.2", "109.1", "572.01", "571.35"],
    ["8v4", "37.15", "36.25", "57.0", "58.7", "558.39", "556.37"],
]

SAFETY_RESULTS = [
    ["二维受控交叉", "0.039", "1743.043", "15.6", "0.0", "基础机制验证，显示解析项可显著抬升最小间距。"],
    ["六自由度4v1危险带", "0.785", "33.638", "19.1", "0.0", "在实际主仿真器中，解析项可将长期危险带占用压到0。"],
    ["六追三多对多交叉", "0.995", "83.723", "1.5", "0.0", "多对多场景下无解析项出现近乎碰撞的换道交叉，有解析项后退出危险区。"],
]

RUNTIME_RESULTS = [
    ["3v1", "0.294", "2.121", "7.203x"],
    ["3v3", "0.106", "1.603", "15.136x"],
    ["5v3", "0.509", "3.488", "6.851x"],
    ["6v3", "0.162", "0.897", "5.550x"],
    ["8v4", "0.714", "1.540", "2.158x"],
]

PLAN_ROWS = [
    ["4月中旬", "补齐Q-learning/传统AC类基线的训练口径、评价指标和参数设置，优先解决其在统一场景下未充分收敛的问题。"],
    ["4月下旬", "继续完成通信丢失条件下的鲁棒性分析，补充更强失联模式和性能退化边界的理论与实验。"],
    ["5月上旬", "围绕平滑因子、解析项强度和危险阈值开展参数敏感性与消融实验，梳理方法边界。"],
    ["5月中旬", "完成毕业论文主体撰写，统一方法描述、公式符号、图表口径和实验叙事。"],
    ["5月下旬—答辩前", "补充最终对比实验、完善答辩材料、整理代码与实验文档，完成论文定稿。"],
]

REFERENCES = [
    "[1] R. Isaacs, Differential Games. New York: Dover, 1999.",
    "[2] V. G. Lopez, G. Bejarano, A. R. Mesquita, and J. P. Hespanha, “Solutions for multiagent pursuit-evasion games on communication graphs: Finite-time capture and asymptotic behaviors,” IEEE Transactions on Automatic Control, vol. 65, no. 5, pp. 1911-1923, 2020.",
    "[3] Z. Xu, D. Yu, Y.-J. Liu, and Z. Wang, “Approximate optimal strategy for multiagent system pursuit-evasion game,” IEEE Systems Journal, vol. 18, no. 3, pp. 1669-1680, 2024.",
    "[4] R. Lowe, Y. Wu, A. Tamar, J. Harb, P. Abbeel, and I. Mordatch, “Multi-agent actor-critic for mixed cooperative-competitive environments,” in Proc. NeurIPS, 2017, pp. 6379-6390.",
    "[5] F. L. Lewis, S. Jagannathan, and A. Yesildirak, Neural Network Control of Robot Manipulators and Nonlinear Systems. London: Taylor & Francis, 1998.",
    "[6] W. Ren and R. W. Beard, “Consensus seeking in multiagent systems under dynamically changing interaction topologies,” IEEE Transactions on Automatic Control, vol. 50, no. 5, pp. 655-661, 2005.",
    "[7] H. Xiong and Y. Zhang, “Reinforcement learning-based formation-surrounding control for multiple quadrotor UAVs pursuit-evasion games,” ISA Transactions, vol. 145, pp. 205-224, 2024.",
    "[8] A. Y. Ng, D. Harada, and S. Russell, “Policy invariance under reward transformations: Theory and application to reward shaping,” in Proc. ICML, 1999, pp. 278-287.",
]

IMAGES = {
    "traj_6v3": ROOT / "paper_figures" / "fig_trajectory_6v3.png",
    "ctrl_team_6v3": ROOT / "paper_figures" / "fig_6v3_ctrl_team.png",
    "collision_traj": ROOT / "paper_figures" / "fig_collision_comparison_traj.png",
    "collision_dmin": ROOT / "paper_figures" / "fig_collision_comparison_dmin.png",
    "full6d_dmin": ROOT / "outputs" / "collision_demo_rigorous_final" / "full_6d" / "fig_6d_dmin_compare.png",
    "many_compare": ROOT / "outputs" / "many_to_many_collision_6v3_final" / "fig_xy_compare.png",
    "many_dmin": ROOT / "outputs" / "many_to_many_collision_6v3_final" / "fig_dmin_compare.png",
    "runtime_compare": ROOT / "outputs" / "full_comm_compare_final" / "6v3" / "compare" / "fig_runtime_compare.png",
}


def clear_paragraph(paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def delete_paragraph(paragraph) -> None:
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def set_run_font(run, size=12, bold=False):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def write_paragraph(paragraph, text: str, size=12, bold=False, align=None):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Cm(0.74) if paragraph.style.name == "Normal" else Cm(0)


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph(style=f"Heading {level}")
    write_paragraph(p, text, size=14 if level == 1 else 12, bold=True)
    p.paragraph_format.first_line_indent = Cm(0)
    return p


def add_normal(doc: Document, text: str):
    p = doc.add_paragraph(style="Normal")
    write_paragraph(p, text)
    return p


def add_formula(doc: Document, text: str):
    p = doc.add_paragraph(style="Normal")
    write_paragraph(p, text, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.first_line_indent = Cm(0)
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph(style="Normal")
    write_paragraph(p, text, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.first_line_indent = Cm(0)


def add_image(doc: Document, image_path: Path, caption: str, width_cm=14.0):
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]]):
    add_caption(doc, title)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, size=10)
    doc.add_paragraph()


def replace_cover(doc: Document):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("学    号"):
            write_paragraph(paragraph, f"学    号   {STUDENT_ID}")
        elif text == "题目":
            write_paragraph(paragraph, TITLE, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        elif text == "年 月 日":
            write_paragraph(paragraph, DATE_TEXT, align=WD_ALIGN_PARAGRAPH.CENTER)

    info_table = doc.tables[0]
    info_table.rows[0].cells[1].text = SCHOOL
    info_table.rows[1].cells[1].text = MAJOR
    info_table.rows[2].cells[1].text = STUDENT
    info_table.rows[3].cells[1].text = ADVISOR
    for row in info_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=12)


def remove_template_body(doc: Document):
    start_idx = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "课题研究背景和目标":
            start_idx = idx
            break
    if start_idx is None:
        return
    for paragraph in list(doc.paragraphs[start_idx:]):
        delete_paragraph(paragraph)


def build_docx():
    doc = Document(str(TEMPLATE))
    replace_cover(doc)
    remove_template_body(doc)

    add_heading(doc, "课题研究背景和目标", 1)
    add_heading(doc, "课题背景", 2)
    add_normal(
        doc,
        "多智能体追逃博弈问题是微分博弈、最优控制与多智能体协同控制交叉形成的典型研究方向。在该类问题中，追逐者一方需要在动态环境下最小化与目标逃避者之间的相对误差并实现有效拦截，逃避者一方则通过最坏情况策略尽可能延迟或规避被捕获。相较于单追单逃问题，多追多场景中的状态维数更高、耦合关系更复杂，同时还叠加了目标分配、组内协同与通信拓扑变化等问题，因此在理论推导和工程实现上都具有明显更高的难度。"
    )
    add_normal(
        doc,
        "结合开题阶段的调研，当前多智能体追逃研究面临的共性困难主要包括四类：一是团队协调性不足，固定目标拓扑难以适应博弈过程中持续变化的几何关系；二是非线性Hamilton-Jacobi-Isaacs方程通常难以解析求解，求解规模又会随着智能体数量增长迅速膨胀；三是工程系统中执行器普遍具有饱和约束，传统二次型代价函数和无约束控制律往往不能直接使用；四是传统Actor-Critic类结构在多智能体场景下需要维护较多网络，在线计算负担较大。"
    )
    add_normal(
        doc,
        "本课题即围绕上述问题展开，研究对象是具有六自由度非线性动力学和控制饱和约束的多追逐者—多逃避者系统。研究起点是完成基线多智能体追逃控制框架的实现，在此基础上进一步加入通信增强与防碰撞扩展，使追逐者在保持原有追踪能力的同时改善组内的空间分离能力和协同效果。"
    )

    add_heading(doc, "课题目标", 2)
    add_normal(
        doc,
        "本课题当前的总体目标可以分为三个层面。第一，完成一个可复现实验链路的多智能体追逃基础框架，包括六自由度动力学模型、动态目标分配、V-SNAC单网络值函数训练、标准场景评估与结果可视化。第二，在基础框架之上设计并实现通信解析项和平滑防碰撞因子，探索其在多追多场景中对组内危险交叉和危险带占用的抑制作用。第三，在保证表述真实的前提下，逐步补齐传统Q-learning/AC类方法的对比实验与通信丢失鲁棒性分析，为毕业论文和后续答辩准备更加完整的理论与实验支撑。"
    )

    add_heading(doc, "课题研究内容", 1)
    add_heading(doc, "多智能体追逃基础框架实现", 2)
    add_normal(
        doc,
        "基础框架的第一部分是状态与动力学建模。当前实现采用六自由度飞行器状态 x=[p_x,p_y,h,v_x,v_y,v_h]^T，并对追逐者和逃避者统一使用控制仿射非线性系统描述。为后续最优控制推导和算法实现服务，代码中已将场景生成、状态初始化、数值积分、捕获判据、目标分配与实验配置统一封装，能够支持3v1、3v3、5v3、6v3、8v4等多种规模场景。"
    )
    add_normal(
        doc,
        "基础框架的第二部分是基于V-SNAC的近似最优控制实现。与传统Actor-Critic不同，当前方法中每个追逐者只维护一个critic参数向量，而不再显式训练独立actor网络；追逐者控制律通过值函数梯度解析给出，并自然嵌入tanh型饱和结构。这样做的直接好处是网络结构轻量、计算链路清晰，也更适合后续加入解析项后进行结构化扩展。"
    )
    add_formula(doc, "u_j^{p*} = -ū_p tanh( (1/(2ū_p)) R_1^{-1} g_j^T ∇V_j(·) )")
    add_normal(
        doc,
        "基础框架的第三部分是动态目标分配。考虑到固定追逐关系在多追多场景下可能导致团队误差增大或局部拥挤，当前系统实现了基于pairwise交换的动态目标分配机制，用于在保持整体稳定的同时改善团队层面的几何关系和追逐效率。"
    )

    add_heading(doc, "通信解析项与平滑防碰撞因子", 2)
    add_normal(
        doc,
        "在基础框架跑通后，本阶段的核心扩展工作是引入通信解析项。该项的出发点不是把“编队”当作独立目标强行叠加，而是针对实际多追多场景中经常出现的组内轨迹重叠、危险交叉和近距离穿越现象，构造一个具有明确结构意义的软分离机制。当前只在同组追逐者之间建立通信边，即仅对同时追逐同一逃避者的追逐者对计算解析协调势。"
    )
    add_normal(
        doc,
        "当前使用的解析项同时包含相对速度一致性项和位置—速度耦合项，并在近距离区域引入平滑pairwise放大因子。当组内两架追逐者接近危险阈值时，平滑因子会连续增大，使解析项在控制律中的影响增强，从而更早地将控制分量分配到空间分离方向，而不是等到非常靠近时才产生明显反应。"
    )
    add_formula(doc, "ρ_jk(x) = d_safe^2 / (||Δp_jk||^2 + ε_d),    α_jk(x) = 1 + 0.5[(ρ_jk-1)+sqrt((ρ_jk-1)^2+ε_α)]")
    add_formula(doc, "Φ_j = (γ/d_j) Σ_k a_jk α_jk(x)[ 0.5||Δv_jk||^2 + ((Δp_jk-Δ^p_jk)^T v_j)/d_ref ]")
    add_normal(
        doc,
        "这里选择平滑因子而不是硬max因子的主要考虑有两点：第一，它在阈值附近连续可微，更利于推导、数值计算和训练稳定性；第二，它在语义上仍然保留“越接近危险距离，协调作用越强”的直观含义，因此比单纯的常数增益更符合当前问题。需要说明的是，这一因子并不意味着绝对不碰撞，而是将其定位为软安全增益。"
    )

    add_heading(doc, "理论推导与当前结论边界", 2)
    add_normal(
        doc,
        "围绕当前方法，已整理出一套较完整的推导链路，主要包括：局部误差动力学表达、带饱和约束的代价泛函、HJI对应的近似最优控制律、V-SNAC单网络值函数近似、以及引入解析项后的增广值函数与控制律表达。对中期报告而言，最重要的不是把所有公式写到极致复杂，而是把“为什么这样设计”讲清楚。"
    )
    add_normal(
        doc,
        "第一，关于V-SNAC结构，当前推导和代码实现表明其最大的优势在于不用为每个智能体额外维护actor网络。值函数采用六维基函数近似，控制律由梯度解析给出，因而在网络规模、参数量和在线控制阶段的计算复杂度上都明显轻于传统AC结构。基础实验中的长训练结果也表明，3v1场景下权重更新能够在给定阈值附近稳定收敛，而不是持续振荡。"
    )
    add_formula(doc, "V̂_j(ẋ_j) = Ŵ_j^T ψ(ẋ_j),    ψ(ẋ_j) ∈ R^6")
    add_normal(
        doc,
        "第二，关于“引入解析项会不会把网络搞坏”，当前可以比较稳妥地说明：解析项和平滑因子本身都是解析可计算的，不增加新的可学习参数，也不改变当前critic的维数和网络数量。也就是说，它扩展的是控制结构和已知解析项，而不是把原有小规模近似器扩展成更高维、更难训练的深网络。因此，从结构复杂度角度，它不会把当前网络架构放大成传统多网络AC那样。"
    )
    add_normal(
        doc,
        "第三，关于“为什么这个因子有效”，当前的解释是：平滑因子通过距离感知增强组内pairwise项，在近距离区域提升解析项梯度的权重；而解析项又通过输入矩阵映射进入控制律，因此在追逐者控制中体现为一种提前介入的软分离驱动力。基础2D受控交叉实验、六自由度危险带实验和六追三多对多交叉实验都支持这一点。"
    )
    add_normal(
        doc,
        "需要如实说明的是，目前这些推导更稳地支撑的是“结构合理、实验有效、网络规模不被放大”这几个结论，而不是“对任意场景的绝对安全保证”或“所有性能指标都单调变好”。尤其在通信丢失、切换拓扑和更复杂多对多交互下，鲁棒性理论仍需继续完善。"
    )

    add_heading(doc, "已完成工作", 1)
    add_heading(doc, "基础系统与工程链路已基本完成", 2)
    add_normal(
        doc,
        "截至目前，基础系统的主体实现已经基本完成。当前代码已经能够完成从训练、评估、作图到报告输出的完整流程，包括标准规模场景训练、多种通信模式评估、传统方法时间对比、碰撞构造场景验证、调试日志输出和轨迹动画生成。为了后续论文整理和问题定位，还额外实现了调试入口、逐步状态日志、训练收敛日志以及面向单场景的live demo脚本。"
    )
    add_normal(
        doc,
        "从项目推进阶段上看，中期前最关键的一步其实不是继续增加新的概念，而是把整个基础框架跑通。当前这一目标已经实现：无论是3v1基础场景、3v3控制对照，还是更大规模的5v3、6v3、8v4场景，都已经能够稳定训练、评估、绘图和输出统计指标。因此，可以认为基础实验平台和主要实验链路已经基本完成。"
    )

    add_heading(doc, "标准场景实验已完成，并得到阶段性结论", 2)
    add_normal(
        doc,
        "围绕主方法，当前已完成3v1、3v3、5v3、6v3、8v4五组标准场景的全量实验，并对全通信、无解析项和通信dropout模式进行了统一统计。整体结果说明：引入解析项后，在若干场景中最小组内间距得到明显改善，且在部分中等规模场景中能够在不明显牺牲捕获时间的情况下改善空间分离；但在较大规模场景中，这种收益并不是对所有指标都单调成立。"
    )
    add_table(
        doc,
        "表1 标准主实验结果摘要（full communication 与 no communication）",
        ["场景", "捕获时间full(s)", "捕获时间no(s)", "d_min full(m)", "d_min no(m)", "均值误差full", "均值误差no"],
        STD_RESULTS,
    )
    add_image(
        doc,
        IMAGES["traj_6v3"],
        "图1 六追三标准场景轨迹图。该图主要用于说明当前多对多基础框架已经能够稳定完成训练、目标切换、评估与可视化流程。",
    )
    add_image(
        doc,
        IMAGES["ctrl_team_6v3"],
        "图2 六追三场景下的控制输入、团队误差与通信dropout对比。该图对应当前会议论文中的主实验组织方式，可用来支撑“基础实验基本完成、通信扩展已可运行”的阶段性判断。",
    )

    add_heading(doc, "碰撞抑制实验已形成从基础到多对多的验证链路", 2)
    add_normal(
        doc,
        "除了标准主实验，本阶段还专门构造了三类碰撞/危险带验证实验。第一类是二维受控交叉实验，用于隔离机制本身：在一个简化但可控的交叉几何中，无解析项时两名追逐者路径发生实际穿越，而引入平滑解析项后，最小间距显著抬升。第二类是六自由度4v1危险带实验，用于在主仿真器中验证解析项确实能把长期危险带占用压到0。第三类是六追三多对多交叉应力场景，用于体现真实多对多配置下，无解析项会出现组内近乎碰撞的换道交叉，而平滑解析项能够将其抬出危险带。"
    )
    add_table(
        doc,
        "表2 碰撞/危险带验证结果摘要",
        ["场景", "min d_min 无解析项(m)", "min d_min 平滑解析项(m)", "危险带占用 无(s)", "危险带占用 有(s)", "说明"],
        SAFETY_RESULTS,
    )
    add_image(
        doc,
        IMAGES["collision_traj"],
        "图3 二维受控交叉场景轨迹对比。该图体现了“无解析项发生物理交叉、有解析项提前横向分离”的基础机制。",
    )
    add_image(
        doc,
        IMAGES["collision_dmin"],
        "图4 二维受控交叉场景的最小间距曲线。该图用于说明平滑解析项能够显著抬升交叉时刻的组内最小间距。",
    )
    add_image(
        doc,
        IMAGES["full6d_dmin"],
        "图5 六自由度危险带实验的最小间距对比。与二维机制图不同，该图是在实际主仿真器中验证解析项的软安全作用。",
    )
    add_image(
        doc,
        IMAGES["many_compare"],
        "图6 六追三多对多交叉应力场景的顶视图对比。该图说明在同一逃避者对应两名追逐者时，无解析项更容易出现近距离换道交叉。",
    )
    add_image(
        doc,
        IMAGES["many_dmin"],
        "图7 六追三多对多交叉应力场景的最小间距对比。该图是当前多对多防碰撞扩展最直接的证据之一。",
    )

    add_heading(doc, "Q-learning/传统AC对比已启动，但尚未收口", 2)
    add_normal(
        doc,
        "围绕传统Q-learning/AC类方法的对比，本阶段已经完成了初步实现与全通信时间口径的对比，能够支持一个比较清楚的阶段性观察：在当前多智能体规模下，V-SNAC单网络结构在在线控制阶段的时间开销明显低于传统AC/Q-learning型结构。"
    )
    add_table(
        doc,
        "表3 V-SNAC与传统AC/Q-learning型基线的单步在线时间对比（初步结果）",
        ["场景", "V-SNAC(ms/step)", "传统AC(ms/step)", "比值(AC/V-SNAC)"],
        RUNTIME_RESULTS,
    )
    add_image(
        doc,
        IMAGES["runtime_compare"],
        "图8 六追三场景下V-SNAC与传统AC/Q-learning型方法的在线时间对比示意。",
    )
    add_normal(
        doc,
        "但这部分工作在中期阶段还不能写成“已经完成最终SOTA对比”。主要原因有两点：其一，当前Q-learning/传统AC类实现虽然已经跑通，但在统一场景、统一预算下尚未完全收敛，不能直接作为最终公平性能结论；其二，对比实验还需要进一步统一训练轮数、超参数和评价口径。因此，目前更适合如实表述为“对比实验正在进行，时间开销差异已经较明确，但最终性能对比仍需继续完善”。"
    )

    add_heading(doc, "阶段性成果", 2)
    add_normal(
        doc,
        "围绕上述工作，当前已形成较系统的代码实现、推导稿、实验报告、调试demo和图像材料。基于本阶段成果，已撰写并投稿一篇会议论文，目前已收到录用结果。中期报告中仅将其作为阶段性研究进展说明，不夸大其结论覆盖范围。整体上看，当前最有把握的结论是：基础框架已实现并基本验证，通信增强与平滑碰撞因子在构造场景中能够发挥作用，而进一步的鲁棒性理论和更规范的SOTA对比仍需继续完成。"
    )

    add_heading(doc, "存在的问题", 1)
    add_heading(doc, "Q-learning/传统方法对比尚未完全完成", 2)
    add_normal(
        doc,
        "当前已经完成传统AC/Q-learning类方法的初步实现和时间计算口径对比，但严格意义上的SOTA性能对比还没有完全收口。一方面，该类基线在当前统一实验配置下尚未完全收敛；另一方面，仍需进一步统一训练轮数、探索设置和评价指标，才能得出足够公平、可用于毕业论文定稿的结论。"
    )
    add_heading(doc, "通信丢失鲁棒性仍需进一步证明", 2)
    add_normal(
        doc,
        "当前已完成15%逐边dropout的实验观察，结果表现为平滑退化，没有出现完全失效。但这仍主要是实验层面的证据，对于更强失联模式、持续通信中断以及动态拓扑切换情况下的性能上界，目前尚缺乏更完整的理论推导。后续需要在此基础上继续完善鲁棒性分析，而不是只停留在经验图像层面。"
    )
    add_heading(doc, "方法边界和参数依赖仍需梳理", 2)
    add_normal(
        doc,
        "从当前实验结果来看，解析项与平滑因子并不是在所有规模、所有指标上都带来单调提升。某些大规模场景中，它更多体现为“改善部分安全指标、但也可能带来一定跟踪代价”的折中特征。因此，后续还需要通过更多参数敏感性和消融实验，把方法的适用条件、收益边界和潜在代价讲清楚。"
    )

    add_heading(doc, "下阶段进度安排", 1)
    add_table(doc, "表4 下阶段工作计划", ["时间", "工作安排"], PLAN_ROWS)

    add_heading(doc, "参考文献", 1)
    for ref in REFERENCES:
        add_normal(doc, ref)

    doc.save(str(DOCX_OUT))


def build_markdown():
    md = f"""# 毕业设计中期报告草稿

- 题目：{TITLE}
- 学号：{STUDENT_ID}
- 学生姓名：{STUDENT}
- 专业：{MAJOR}
- 学院：{SCHOOL}
- 指导教师：{ADVISOR}
- 日期：{DATE_TEXT}

正式 Word 版见：
- `{DOCX_OUT}`

答辩/中检问答准备见：
- `{QA_OUT}`
"""
    MD_OUT.write_text(md, encoding="utf-8")


def build_questions():
    text = """# 中期检查可能问题与回答思路

## 1. 你现在的工作重心和开题时相比有什么变化？
当前已经从“以复现为主”转为“在完整基础框架上做自己的扩展”。基础追逃控制、动态目标分配和V-SNAC训练链路已经跑通，接下来重点集中在通信增强、防碰撞因子、Q-learning对比和通信鲁棒性分析上。

## 2. 目前最确定已经完成的工作是什么？
最确定的是整个基础框架已经实现并完成主实验，包括六自由度动力学、V-SNAC训练、标准规模场景评估、调试demo、图像与动画输出。其次，解析通信项和平滑因子已经在二维受控交叉、六自由度危险带和六追三多对多交叉场景中显示出明确作用。

## 3. 平滑因子为什么有效？
因为它在追逐者近距离时会连续放大组内pairwise解析项，使更多控制分量提前分配到空间分离方向，从而降低近距离交叉和危险带占用。它是软安全增益，不是硬约束屏障。

## 4. 平滑因子会不会让网络训练更差？
从当前结构上看，它不会增加新的可学习参数，也不增加critic的维数和网络数量，因此不会把原来的小规模V-SNAC架构放大成更难训练的大网络。但它会改变控制结构和部分实验表现，所以更准确的说法是“不会扩大网络规模，不保证所有性能指标都更优”。

## 5. 为什么你还不能说实现了绝对防碰撞？
因为当前方法是解析软分离机制，不是硬约束安全控制。现有结果只能说明在构造场景和部分标准场景中，解析项显著改善了组内最小间距和危险带占用时间，还不能推出任意初值和任意场景下的绝对不碰撞结论。

## 6. 现在Q-learning对比做到什么程度？
已经完成了传统AC/Q-learning类基线的初步实现和时间计算口径对比，说明当前V-SNAC结构在在线控制阶段更轻。但严格、公平的SOTA性能对比还没有完成，因为该类基线在统一配置下尚未完全收敛，参数也还需要继续校准。

## 7. 通信鲁棒性现在能说到什么程度？
可以说已经完成了15%逐边dropout的实验验证，结果表现为平滑退化，没有立刻失效；但更强失联模式和理论上界还没有完成，因此后续还需要继续补充。

## 8. 为什么要做多个碰撞场景，而不是只做一个？
因为二维受控交叉更适合展示机制本身，六自由度危险带实验更适合说明在主仿真器中该机制能否真正生效，六追三多对多交叉则更贴近本课题真正关注的多追多场景。三者组合起来，比单一图像更有说服力。

## 9. 你现在最需要补的两块是什么？
第一是Q-learning/传统AC类SOTA对比的最终收口，第二是通信丢失和动态拓扑下鲁棒性与稳定性分析的进一步完善。

## 10. 如果老师问你中期最客观的结论是什么？
最客观的结论是：基础框架实现和主实验已经基本完成；引入的平滑解析通信项在构造碰撞场景和部分标准场景中能够发挥作用；但最终的SOTA对比和更强鲁棒性证明仍在进行，毕业论文阶段还需要继续补全。
"""
    QA_OUT.write_text(text, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_markdown()
    build_questions()
    print(DOCX_OUT)
    print(MD_OUT)
    print(QA_OUT)
